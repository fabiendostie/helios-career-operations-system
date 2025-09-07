"""
File ingestion and language detection components.
"""

import logging
import json
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Optional, Any

import yaml
from pypdf import PdfReader
from docx import Document as DocxDocument
import mistune
from langdetect import detect, DetectorFactory

from ..schemas.master_schema import Document


# Set langdetect seed for consistent results
DetectorFactory.seed = 0


class FileHandler(ABC):
    """Abstract base class for file handlers using Strategy pattern."""

    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        """Extract text content from file."""
        pass


class PDFHandler(FileHandler):
    """Handle PDF file text extraction."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from PDF using pypdf."""
        try:
            reader = PdfReader(str(file_path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logging.error(f"Error reading PDF {file_path}: {e}")
            return ""


class DOCXHandler(FileHandler):
    """Handle DOCX file text extraction."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            doc = DocxDocument(str(file_path))
            text = ""

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text.strip()
        except Exception as e:
            logging.error(f"Error reading DOCX {file_path}: {e}")
            return ""


class MarkdownHandler(FileHandler):
    """Handle Markdown file text extraction."""

    def __init__(self):
        self.renderer = mistune.create_markdown(renderer="ast")

    def extract_text(self, file_path: Path) -> str:
        """Extract text from Markdown using mistune."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Parse markdown to AST then extract text
            ast = self.renderer(content)
            return self._extract_text_from_ast(ast)
        except UnicodeDecodeError:
            # Try with latin-1 encoding
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
                ast = self.renderer(content)
                return self._extract_text_from_ast(ast)
            except Exception as e:
                logging.error(f"Error reading Markdown {file_path}: {e}")
                return ""
        except Exception as e:
            logging.error(f"Error reading Markdown {file_path}: {e}")
            return ""

    def _extract_text_from_ast(self, node) -> str:
        """Recursively extract text from markdown AST."""
        if isinstance(node, dict):
            if node.get("type") == "text":
                return node.get("raw", "")
            elif "children" in node:
                return "".join(
                    self._extract_text_from_ast(child) for child in node["children"]
                )
        elif isinstance(node, list):
            return "".join(self._extract_text_from_ast(item) for item in node)

        return str(node) if node else ""


class TextHandler(FileHandler):
    """Handle plain text file extraction."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        encodings = ["utf-8", "latin-1", "cp1252", "utf-16"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                # Remove BOM if present
                if content.startswith("\ufeff"):
                    content = content[1:]
                return content.strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logging.error(f"Error reading text file {file_path}: {e}")
                return ""

        logging.error(f"Could not decode text file {file_path} with any encoding")
        return ""


class YAMLHandler(FileHandler):
    """Handle YAML file text extraction."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from YAML file by converting to readable format."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)

            return self._yaml_to_text(data)
        except Exception as e:
            logging.error(f"Error reading YAML {file_path}: {e}")
            return ""

    def _yaml_to_text(self, data: Any, indent: int = 0) -> str:
        """Convert YAML data structure to readable text."""
        if isinstance(data, dict):
            text = ""
            for key, value in data.items():
                text += "  " * indent + f"{key}:\n"
                if isinstance(value, (dict, list)):
                    text += self._yaml_to_text(value, indent + 1)
                else:
                    text += "  " * (indent + 1) + f"{value}\n"
            return text
        elif isinstance(data, list):
            text = ""
            for item in data:
                if isinstance(item, (dict, list)):
                    text += self._yaml_to_text(item, indent)
                else:
                    text += "  " * indent + f"- {item}\n"
            return text
        else:
            return str(data)


class JSONHandler(FileHandler):
    """Handle JSON file text extraction."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from JSON file by flattening structure."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            return self._json_to_text(data)
        except Exception as e:
            logging.error(f"Error reading JSON {file_path}: {e}")
            return ""

    def _json_to_text(self, data: Any, prefix: str = "") -> str:
        """Flatten JSON structure to readable text."""
        if isinstance(data, dict):
            text = ""
            for key, value in data.items():
                current_key = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    text += f"{current_key}:\n"
                    text += self._json_to_text(value, current_key)
                else:
                    text += f"{current_key}: {value}\n"
            return text
        elif isinstance(data, list):
            text = ""
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    text += self._json_to_text(item, f"{prefix}[{i}]")
                else:
                    text += f"{prefix}[{i}]: {item}\n"
            return text
        else:
            return str(data)


class LanguageDetector:
    """Detect language of text content."""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def detect_language(self, text: str) -> str:
        """
        Detect language of text content.

        Args:
            text: Input text to analyze

        Returns:
            'en' for English, 'fr' for French, 'en' as default
        """
        if not text or len(text.strip()) < 10:
            self.logger.warning(
                "Text too short for reliable language detection, defaulting to English"
            )
            return "en"

        try:
            detected = detect(text)
            # Map language codes to our supported languages
            if detected in ["en", "english"]:
                return "en"
            elif detected in ["fr", "french"]:
                return "fr"
            else:
                self.logger.warning(
                    f"Detected language '{detected}' not supported, defaulting to English"
                )
                return "en"
        except Exception as e:
            self.logger.error(f"Language detection failed: {e}, defaulting to English")
            return "en"


class IngestionEngine:
    """Main engine for file ingestion and processing."""

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.language_detector = LanguageDetector()

        # File handlers using Strategy pattern
        self.file_handlers: Dict[str, FileHandler] = {
            ".pdf": PDFHandler(),
            ".docx": DOCXHandler(),
            ".md": MarkdownHandler(),
            ".txt": TextHandler(),
            ".yml": YAMLHandler(),
            ".yaml": YAMLHandler(),
            ".json": JSONHandler(),
        }

    def ingest_files(self, directory_path: Path) -> List[Document]:
        """
        Read all supported files from directory and detect languages.

        Args:
            directory_path: Path to directory containing resume files

        Returns:
            List of Document objects with content and metadata
        """
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        if not directory_path.is_dir():
            raise ValueError(f"Path is not a directory: {directory_path}")

        documents = []
        supported_extensions = self.file_handlers.keys()

        # Find all supported files
        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    document = self._read_file(file_path)
                    if document:
                        documents.append(document)
                        self.logger.info(
                            f"Processed {file_path.name}: {len(document.content)} chars, language: {document.language}"
                        )
                    else:
                        self.logger.warning(f"Failed to process {file_path}")
                except Exception as e:
                    self.logger.error(f"Error processing {file_path}: {e}")
                    continue

        self.logger.info(
            f"Successfully processed {len(documents)} documents from {directory_path}"
        )
        return documents

    def _read_file(self, file_path: Path) -> Optional[Document]:
        """
        Read and process a single file using appropriate handler.

        Args:
            file_path: Path to the file to process

        Returns:
            Document object or None if processing failed
        """
        file_extension = file_path.suffix.lower()

        if file_extension not in self.file_handlers:
            self.logger.warning(
                f"Unsupported file type: {file_extension} for {file_path}"
            )
            return None

        try:
            handler = self.file_handlers[file_extension]
            content = handler.extract_text(file_path)

            if not content.strip():
                self.logger.warning(f"No content extracted from {file_path}")
                return None

            # Detect language
            language = self.language_detector.detect_language(content)

            # Create metadata
            metadata = {
                "file_size": file_path.stat().st_size,
                "modified_time": file_path.stat().st_mtime,
            }

            return Document(
                file_path=file_path,
                content=content,
                language=language,
                file_type=file_extension,
                metadata=metadata,
            )

        except Exception as e:
            self.logger.error(f"Failed to read file {file_path}: {e}")
            return None
