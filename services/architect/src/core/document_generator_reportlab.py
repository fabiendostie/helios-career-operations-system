"""Multi-format document generation with ReportLab PDF engine."""

import os
import gc
import time
import tempfile
import asyncio
from typing import Dict, Any, Optional, List, Union
from pathlib import Path
import mimetypes
from contextlib import asynccontextmanager

import structlog
import psutil
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

from .config import get_settings, PERFORMANCE_THRESHOLDS
from .template_engine import TemplateEngine, TemplateRenderError

logger = structlog.get_logger(__name__)


class DocumentGenerationError(Exception):
    """Raised when document generation fails."""
    pass


class ReportLabPDFGenerator:
    """PDF generation using ReportLab for maximum compatibility."""

    def __init__(self):
        self.settings = get_settings()
        self.styles = getSampleStyleSheet()
        self._create_custom_styles()

    def _create_custom_styles(self):
        """Create custom styles for ATS-compliant PDF generation."""
        # Professional heading style
        self.styles.add(ParagraphStyle(
            name='ATSHeading1',
            parent=self.styles['Heading1'],
            fontSize=16,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            spaceAfter=12,
            textColor='black'
        ))

        self.styles.add(ParagraphStyle(
            name='ATSHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            fontName='Helvetica-Bold',
            alignment=TA_LEFT,
            spaceBefore=12,
            spaceAfter=6,
            textColor='black'
        ))

        # Body text style
        self.styles.add(ParagraphStyle(
            name='ATSBody',
            parent=self.styles['Normal'],
            fontSize=11,
            fontName='Helvetica',
            alignment=TA_LEFT,
            spaceBefore=3,
            spaceAfter=3,
            leading=13
        ))

        # Contact info style
        self.styles.add(ParagraphStyle(
            name='ContactInfo',
            parent=self.styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            alignment=TA_CENTER,
            spaceBefore=6,
            spaceAfter=12
        ))

        # Skills style
        self.styles.add(ParagraphStyle(
            name='Skills',
            parent=self.styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            alignment=TA_LEFT,
            spaceBefore=6,
            spaceAfter=6
        ))

    async def generate_pdf(self, resume_data: Dict[str, Any], template_type: str = "classic") -> bytes:
        """Generate PDF directly from structured resume data."""
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_path = temp_file.name

            # Create PDF document
            doc = SimpleDocTemplate(
                temp_path,
                pagesize=letter,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch
            )

            story = []

            # Header section
            story.extend(self._create_header_section(resume_data))
            story.append(Spacer(1, 12))

            # Impact Summary (T-shaped top bar)
            story.extend(self._create_impact_summary(resume_data))
            story.append(Spacer(1, 12))

            # Experience section (T-shaped vertical bar)
            story.extend(self._create_experience_section(resume_data))
            story.append(Spacer(1, 12))

            # Education section
            if resume_data.get('education'):
                story.extend(self._create_education_section(resume_data))
                story.append(Spacer(1, 6))

            # Skills section
            if resume_data.get('core_skills'):
                story.extend(self._create_skills_section(resume_data))

            # Build PDF
            await asyncio.to_thread(doc.build, story)

            # Read generated PDF
            with open(temp_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()

            # Cleanup
            try:
                os.unlink(temp_path)
            except:
                pass

            return pdf_content

        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise DocumentGenerationError(f"PDF generation failed: {e}")

    def _create_header_section(self, resume_data: Dict[str, Any]) -> List:
        """Create header section with name and contact info."""
        elements = []

        # Candidate name
        name = resume_data.get('candidate_name', 'Professional')
        elements.append(Paragraph(name, self.styles['ATSHeading1']))

        # Contact information
        contact_parts = []
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('location'):
            contact_parts.append(resume_data['location'])

        if contact_parts:
            contact_text = ' | '.join(contact_parts)
            elements.append(Paragraph(contact_text, self.styles['ContactInfo']))

        return elements

    def _create_impact_summary(self, resume_data: Dict[str, Any]) -> List:
        """Create impact summary section (T-shaped top bar)."""
        elements = []

        summary_text = resume_data.get('summary', '')
        if not summary_text and resume_data.get('strategic_metadata'):
            # Generate summary from metadata
            meta = resume_data['strategic_metadata']
            adjective = meta.get('adjective', 'Experienced')
            role = resume_data.get('target_role_title', 'Professional')
            years = meta.get('years_experience', '5+')
            competencies = meta.get('core_competencies', [])[:3]

            if competencies:
                competency_text = ', '.join(competencies[:2])
                if len(competencies) > 2:
                    competency_text += f", and {competencies[2]}"

                summary_text = f"{adjective} {role} with {years} years of experience specializing in {competency_text}."

        if summary_text:
            elements.append(Paragraph(summary_text, self.styles['ATSBody']))
            elements.append(Spacer(1, 6))

        # Core skills
        core_skills = resume_data.get('core_skills', [])
        if core_skills:
            skills_text = f"<b>Core Competencies:</b> {' • '.join(core_skills[:8])}"
            elements.append(Paragraph(skills_text, self.styles['Skills']))

        return elements

    def _create_experience_section(self, resume_data: Dict[str, Any]) -> List:
        """Create experience section (T-shaped vertical bar)."""
        elements = []
        elements.append(Paragraph("Professional Experience", self.styles['ATSHeading2']))

        work_experience = resume_data.get('work_experience', [])
        for job in work_experience:
            # Job header
            job_title = job.get('role_title', '')
            company = job.get('company_name', '')
            duration = job.get('duration', '')

            if job_title and company:
                header_text = f"<b>{job_title}</b> - {company}"
                if duration:
                    header_text += f" ({duration})"
                elements.append(Paragraph(header_text, self.styles['ATSBody']))
                elements.append(Spacer(1, 3))

            # Accomplishments
            accomplishments = job.get('accomplishments', [])
            for accomplishment in accomplishments:
                elements.append(Paragraph(f"• {accomplishment}", self.styles['ATSBody']))

            # Skills used
            skills_used = job.get('skills_used', [])
            if skills_used:
                skills_text = f"<i>Technologies:</i> {', '.join(skills_used)}"
                elements.append(Paragraph(skills_text, self.styles['Skills']))

            elements.append(Spacer(1, 12))

        return elements

    def _create_education_section(self, resume_data: Dict[str, Any]) -> List:
        """Create education section."""
        elements = []
        elements.append(Paragraph("Education", self.styles['ATSHeading2']))

        education = resume_data.get('education', [])
        for edu in education:
            degree = edu.get('degree', '')
            institution = edu.get('institution', '')
            year = edu.get('year', '')

            if degree and institution:
                edu_text = f"<b>{degree}</b> - {institution}"
                if year:
                    edu_text += f" ({year})"
                elements.append(Paragraph(edu_text, self.styles['ATSBody']))

        return elements

    def _create_skills_section(self, resume_data: Dict[str, Any]) -> List:
        """Create additional skills section."""
        elements = []

        skills_inventory = resume_data.get('skills_inventory', {})
        if skills_inventory:
            elements.append(Paragraph("Technical Skills", self.styles['ATSHeading2']))

            for category, skills in skills_inventory.items():
                if isinstance(skills, list) and skills:
                    category_name = category.replace('_', ' ').title()
                    skills_text = f"<b>{category_name}:</b> {', '.join(skills)}"
                    elements.append(Paragraph(skills_text, self.styles['ATSBody']))

        return elements


class MultiFormatDocumentGenerator:
    """Document generator supporting multiple output formats."""

    def __init__(self):
        self.settings = get_settings()
        self.pdf_generator = ReportLabPDFGenerator()
        self.template_engine = TemplateEngine()

    async def generate_document(
        self,
        resume_data: Dict[str, Any],
        output_format: str = "pdf",
        template_type: str = "classic"
    ) -> bytes:
        """Generate document in specified format."""

        start_time = time.time()
        logger.info(f"Starting document generation",
                   format=output_format, template=template_type)

        try:
            if output_format.lower() == "pdf":
                result = await self.pdf_generator.generate_pdf(resume_data, template_type)
            elif output_format.lower() == "docx":
                result = await self._generate_docx(resume_data, template_type)
            elif output_format.lower() == "html":
                result = await self._generate_html(resume_data, template_type)
            elif output_format.lower() == "markdown":
                result = await self._generate_markdown(resume_data, template_type)
            else:
                raise DocumentGenerationError(f"Unsupported format: {output_format}")

            generation_time = time.time() - start_time
            logger.info(f"Document generated successfully",
                       format=output_format,
                       generation_time=generation_time,
                       size_bytes=len(result))

            return result

        except Exception as e:
            logger.error(f"Document generation failed",
                        format=output_format,
                        error=str(e))
            raise

    async def _generate_docx(self, resume_data: Dict[str, Any], template_type: str) -> bytes:
        """Generate DOCX document."""
        doc = Document()

        # Header
        name = resume_data.get('candidate_name', 'Professional')
        header = doc.add_heading(name, 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = []
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('location'):
            contact_parts.append(resume_data['location'])

        if contact_parts:
            contact = doc.add_paragraph(' | '.join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary
        summary_text = resume_data.get('summary', '')
        if summary_text:
            doc.add_paragraph(summary_text)

        # Experience
        if resume_data.get('work_experience'):
            doc.add_heading('Professional Experience', level=1)
            for job in resume_data['work_experience']:
                job_title = job.get('role_title', '')
                company = job.get('company_name', '')
                duration = job.get('duration', '')

                if job_title and company:
                    job_header = f"{job_title} - {company}"
                    if duration:
                        job_header += f" ({duration})"
                    doc.add_paragraph(job_header, style='Heading 2')

                for accomplishment in job.get('accomplishments', []):
                    p = doc.add_paragraph(accomplishment, style='List Bullet')

        # Education
        if resume_data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in resume_data['education']:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                year = edu.get('year', '')

                if degree and institution:
                    edu_text = f"{degree} - {institution}"
                    if year:
                        edu_text += f" ({year})"
                    doc.add_paragraph(edu_text)

        # Save to bytes using BytesIO
        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    async def _generate_html(self, resume_data: Dict[str, Any], template_type: str) -> bytes:
        """Generate HTML document."""
        # Use template engine for HTML generation
        template_path = f"resume/{template_type}.j2"
        html_content = await self.template_engine.render_template(template_path, resume_data)
        return html_content.encode('utf-8')

    async def _generate_markdown(self, resume_data: Dict[str, Any], template_type: str) -> bytes:
        """Generate Markdown document."""
        lines = []

        # Header
        name = resume_data.get('candidate_name', 'Professional')
        lines.append(f"# {name}")
        lines.append("")

        # Contact
        contact_parts = []
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('location'):
            contact_parts.append(resume_data['location'])

        if contact_parts:
            lines.append(' | '.join(contact_parts))
            lines.append("")

        # Summary
        summary_text = resume_data.get('summary', '')
        if summary_text:
            lines.append(summary_text)
            lines.append("")

        # Experience
        if resume_data.get('work_experience'):
            lines.append("## Professional Experience")
            lines.append("")

            for job in resume_data['work_experience']:
                job_title = job.get('role_title', '')
                company = job.get('company_name', '')
                duration = job.get('duration', '')

                if job_title and company:
                    job_header = f"### {job_title} - {company}"
                    if duration:
                        job_header += f" ({duration})"
                    lines.append(job_header)
                    lines.append("")

                for accomplishment in job.get('accomplishments', []):
                    lines.append(f"- {accomplishment}")
                lines.append("")

        # Education
        if resume_data.get('education'):
            lines.append("## Education")
            lines.append("")

            for edu in resume_data['education']:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                year = edu.get('year', '')

                if degree and institution:
                    edu_text = f"**{degree}** - {institution}"
                    if year:
                        edu_text += f" ({year})"
                    lines.append(edu_text)
            lines.append("")

        return '\n'.join(lines).encode('utf-8')
