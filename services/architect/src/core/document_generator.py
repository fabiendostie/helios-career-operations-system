"""Multi-format document generation with memory optimization."""

import os
import gc
import time
import tempfile
import asyncio
from typing import Dict, Any, Optional, List, Union
from pathlib import Path
import mimetypes
from contextlib import asynccontextmanager

import structlog
import psutil
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

from .config import get_settings, PERFORMANCE_THRESHOLDS
from .template_engine import TemplateEngine, TemplateRenderError

logger = structlog.get_logger(__name__)


class DocumentGenerationError(Exception):
    """Raised when document generation fails."""
    pass


class MemoryEfficientPDFGenerator:
    """PDF generation with memory optimization."""
    
    def __init__(self):
        self.settings = get_settings()
        
    @asynccontextmanager
    async def memory_managed_pdf_conversion(self, html_content: str):
        """Context manager for memory-efficient PDF conversion."""
        temp_html_file = None
        temp_pdf_file = None
        
        try:
            # Use temporary files to avoid keeping large content in memory
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as html_file:
                html_file.write(html_content)
                temp_html_file = html_file.name
            
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
                temp_pdf_file = pdf_file.name
            
            # Memory-efficient PDF generation with WeasyPrint
            logger.debug("Starting PDF generation", temp_html=temp_html_file)
            
            # Configure WeasyPrint for memory efficiency and ATS compliance
            html_doc = weasyprint.HTML(filename=temp_html_file)
            
            # Generate PDF with ATS-optimized settings
            start_time = time.time()
            pdf_content = html_doc.write_pdf(
                optimize_size=True,  # Optimize for smaller file size
                pdf_version='1.4',   # Use PDF 1.4 for maximum compatibility
                presentational_hints=True,  # Include CSS hints for better rendering
                # Additional WeasyPrint options for ATS compliance
                page_size='letter',  # Standard US letter size
                margin='0.75in'     # Professional margins
            )
            generation_time = time.time() - start_time
            
            logger.info("PDF generated successfully", 
                       generation_time=generation_time,
                       file_size=len(pdf_content))
            
            yield pdf_content
            
        except MemoryError:
            logger.error("Memory error during PDF generation")
            # Attempt recovery with minimal settings
            yield await self._generate_minimal_pdf(html_content)
            
        except Exception as e:
            logger.error("PDF generation failed", error=str(e), exc_info=True)
            raise DocumentGenerationError(f"PDF generation failed: {e}")
            
        finally:
            # Cleanup temporary files
            for temp_file in [temp_html_file, temp_pdf_file]:
                if temp_file and os.path.exists(temp_file):
                    try:
                        os.unlink(temp_file)
                    except:
                        pass
            
            # Force cleanup
            gc.collect()
    
    async def _generate_minimal_pdf(self, html_content: str) -> bytes:
        """Generate PDF with minimal memory usage."""
        # Simplified HTML content to reduce memory usage
        simplified_html = self._simplify_html_content(html_content)
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html') as html_file:
            html_file.write(simplified_html)
            html_file.flush()
            
            html_doc = weasyprint.HTML(filename=html_file.name)
            return html_doc.write_pdf(optimize_size=True, pdf_version='1.4')
    
    def _simplify_html_content(self, html_content: str) -> str:
        """Simplify HTML content to reduce memory usage."""
        import re
        
        # Remove comments
        html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)
        
        # Compress whitespace
        html_content = re.sub(r'\s+', ' ', html_content)
        
        # Remove empty elements
        html_content = re.sub(r'<(\w+)[^>]*>\s*</\1>', '', html_content)
        
        return html_content.strip()


class DOCXGenerator:
    """DOCX document generation for legacy ATS compatibility."""
    
    def __init__(self):
        self.settings = get_settings()
    
    async def generate_docx_from_data(
        self, 
        career_data: Dict[str, Any], 
        template_style: str = "classic"
    ) -> bytes:
        """Generate DOCX document directly from career data."""
        try:
            # Create new document
            doc = Document()
            
            # Set document margins (ATS-friendly)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add header with candidate name and contact info
            self._add_header_section(doc, career_data)
            
            # Add impact summary (T-shaped top bar)
            self._add_impact_summary(doc, career_data)
            
            # Add professional experience (T-shaped vertical bar)
            self._add_experience_section(doc, career_data)
            
            # Add education section
            if career_data.get('education'):
                self._add_education_section(doc, career_data)
            
            # Add technical skills
            if career_data.get('skills_inventory'):
                self._add_skills_section(doc, career_data)
            
            # Save to bytes
            with tempfile.NamedTemporaryFile() as tmp_file:
                doc.save(tmp_file.name)
                tmp_file.seek(0)
                docx_content = tmp_file.read()
            
            logger.info("DOCX generated successfully", 
                       file_size=len(docx_content))
            
            return docx_content
            
        except Exception as e:
            logger.error("DOCX generation failed", error=str(e), exc_info=True)
            raise DocumentGenerationError(f"DOCX generation failed: {e}")
    
    def _add_header_section(self, doc: Document, career_data: Dict[str, Any]):
        """Add header with name and contact information."""
        # Candidate name
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(career_data.get('candidate_name', 'Professional'))
        name_run.font.size = Pt(16)
        name_run.bold = True
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_info = []
        if career_data.get('email'):
            contact_info.append(career_data['email'])
        if career_data.get('phone'):
            contact_info.append(career_data['phone'])
        if career_data.get('location'):
            contact_info.append(career_data['location'])
            
        if contact_info:
            contact_paragraph = doc.add_paragraph()
            contact_run = contact_paragraph.add_run(' | '.join(contact_info))
            contact_run.font.size = Pt(10)
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add space
        doc.add_paragraph()
    
    def _add_impact_summary(self, doc: Document, career_data: Dict[str, Any]):
        """Add T-shaped impact summary."""
        strategic_metadata = career_data.get('strategic_metadata', {})
        core_competencies = strategic_metadata.get('core_competencies', [])
        achievements = strategic_metadata.get('quantified_achievements', [])
        
        # Create impact summary text
        years_exp = self._calculate_years_experience(career_data.get('work_experience', []))
        summary_text = (
            f"Experienced Professional with {years_exp} years of experience "
            f"specializing in {', '.join(core_competencies[:3])}. "
        )
        
        if achievements:
            summary_text += f"Proven ability to {achievements[0]} through strategic expertise."
        
        summary_paragraph = doc.add_paragraph()
        summary_run = summary_paragraph.add_run(summary_text)
        summary_run.font.size = Pt(11)
        summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Core skills
        skills_inventory = career_data.get('skills_inventory', {})
        core_skills = self._extract_core_skills(skills_inventory)[:8]
        
        if core_skills:
            doc.add_paragraph()  # Space
            skills_paragraph = doc.add_paragraph()
            skills_header = skills_paragraph.add_run("Core Competencies: ")
            skills_header.bold = True
            skills_paragraph.add_run(' • '.join(core_skills))
        
        doc.add_paragraph()  # Space
    
    def _add_experience_section(self, doc: Document, career_data: Dict[str, Any]):
        """Add professional experience section."""
        # Section header
        exp_header = doc.add_paragraph()
        exp_header_run = exp_header.add_run("PROFESSIONAL EXPERIENCE")
        exp_header_run.font.size = Pt(14)
        exp_header_run.bold = True
        
        # Add horizontal line effect with underline
        exp_header_run.underline = True
        
        work_experience = career_data.get('work_experience', [])
        
        for job in work_experience:
            doc.add_paragraph()  # Space
            
            # Job title and company
            job_header = doc.add_paragraph()
            title_run = job_header.add_run(job.get('role_title', 'Professional'))
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            company_text = f" - {job.get('company_name', 'Company')}"
            if job.get('start_date') or job.get('end_date'):
                duration = f"{job.get('start_date', '')} - {job.get('end_date', 'Present')}"
                company_text += f" ({duration})"
            
            job_header.add_run(company_text)
            
            # Accomplishments
            accomplishments = job.get('accomplishments', [])
            for accomplishment in accomplishments:
                bullet_paragraph = doc.add_paragraph(style='List Bullet')
                bullet_paragraph.add_run(accomplishment)
            
            # Skills used
            skills_used = job.get('skills_used', [])
            if skills_used:
                skills_paragraph = doc.add_paragraph()
                tech_run = skills_paragraph.add_run("Technologies: ")
                tech_run.italic = True
                tech_run.font.size = Pt(9)
                skills_paragraph.add_run(', '.join(skills_used))
                skills_paragraph.runs[-1].font.size = Pt(9)
    
    def _add_education_section(self, doc: Document, career_data: Dict[str, Any]):
        """Add education section."""
        doc.add_paragraph()  # Space
        
        # Section header
        edu_header = doc.add_paragraph()
        edu_header_run = edu_header.add_run("EDUCATION")
        edu_header_run.font.size = Pt(14)
        edu_header_run.bold = True
        edu_header_run.underline = True
        
        education = career_data.get('education', [])
        
        for edu in education:
            doc.add_paragraph()  # Space
            
            edu_paragraph = doc.add_paragraph()
            degree_run = edu_paragraph.add_run(edu.get('degree', 'Degree'))
            degree_run.bold = True
            
            if edu.get('institution'):
                edu_paragraph.add_run(f" - {edu['institution']}")
            
            if edu.get('year'):
                edu_paragraph.add_run(f" ({edu['year']})")
    
    def _add_skills_section(self, doc: Document, career_data: Dict[str, Any]):
        """Add technical skills section."""
        doc.add_paragraph()  # Space
        
        # Section header
        skills_header = doc.add_paragraph()
        skills_header_run = skills_header.add_run("TECHNICAL SKILLS")
        skills_header_run.font.size = Pt(14)
        skills_header_run.bold = True
        skills_header_run.underline = True
        
        skills_inventory = career_data.get('skills_inventory', {})
        technical_skills = skills_inventory.get('technical_skills', {})
        
        for category, skills in technical_skills.items():
            if skills:
                doc.add_paragraph()  # Space
                
                category_paragraph = doc.add_paragraph()
                category_run = category_paragraph.add_run(f"{category.capitalize()}: ")
                category_run.bold = True
                
                if isinstance(skills, list):
                    category_paragraph.add_run(', '.join(skills))
                else:
                    category_paragraph.add_run(str(skills))
    
    def _calculate_years_experience(self, work_experience: List[Dict[str, Any]]) -> int:
        """Calculate years of experience."""
        if not work_experience:
            return 0
        return min(len(work_experience) * 2, 20)  # Cap at 20 years
    
    def _extract_core_skills(self, skills_inventory: Dict[str, Any]) -> List[str]:
        """Extract core skills from skills inventory."""
        core_skills = []
        
        technical_skills = skills_inventory.get('technical_skills', {})
        for category, skills in technical_skills.items():
            if isinstance(skills, list):
                core_skills.extend(skills[:3])
        
        soft_skills = skills_inventory.get('soft_skills', [])
        core_skills.extend(soft_skills[:3])
        
        return core_skills[:8]


class DocumentGenerator:
    """High-level document generation interface supporting multiple formats."""
    
    def __init__(self):
        self.settings = get_settings()
        self.template_engine = TemplateEngine()
        self.pdf_generator = MemoryEfficientPDFGenerator()
        self.docx_generator = DOCXGenerator()
        
        logger.info("Document generator initialized")
    
    async def generate_document(
        self,
        template_name: str,
        career_data: Dict[str, Any],
        output_format: str = "pdf",
        job_requirements: Optional[Dict[str, Any]] = None,
        customizations: Optional[Dict[str, Any]] = None,
        document_type: str = "resume"
    ) -> Dict[str, Any]:
        """
        Generate document in specified format.
        
        Args:
            template_name: Name of template to use
            career_data: Career data from Master Career Database
            output_format: Target format (pdf, html, markdown, docx)
            job_requirements: Optional job requirements for customization
            customizations: Optional template customizations
            document_type: Type of document (resume, cover_letter)
            
        Returns:
            Dictionary with document content and metadata
        """
        start_time = time.time()
        initial_memory = psutil.Process().memory_info().rss / 1024 / 1024
        
        try:
            logger.info(
                "Starting document generation",
                template_name=template_name,
                output_format=output_format,
                document_type=document_type
            )
            
            # Generate HTML content first (base format)
            if document_type == "resume":
                html_content = await self.template_engine.render_resume_template(
                    template_name, career_data, job_requirements, customizations
                )
            elif document_type == "cover_letter":
                html_content = await self.template_engine.render_cover_letter_template(
                    template_name, career_data, job_requirements, customizations
                )
            else:
                raise DocumentGenerationError(f"Unsupported document type: {document_type}")
            
            # Convert to requested format
            if output_format.lower() == "html":
                content = html_content.encode('utf-8')
                mime_type = "text/html"
                
            elif output_format.lower() == "pdf":
                async with self.pdf_generator.memory_managed_pdf_conversion(html_content) as pdf_content:
                    content = pdf_content
                mime_type = "application/pdf"
                
            elif output_format.lower() == "markdown":
                # Convert HTML to Markdown (simplified)
                markdown_content = self._html_to_markdown(html_content)
                content = markdown_content.encode('utf-8')
                mime_type = "text/markdown"
                
            elif output_format.lower() == "docx":
                # Generate DOCX directly from career data for better compatibility
                content = await self.docx_generator.generate_docx_from_data(
                    career_data, template_name
                )
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                
            else:
                raise DocumentGenerationError(f"Unsupported output format: {output_format}")
            
            # Calculate performance metrics
            generation_time = time.time() - start_time
            final_memory = psutil.Process().memory_info().rss / 1024 / 1024
            memory_delta = final_memory - initial_memory
            
            # Log performance warnings if needed
            if generation_time > PERFORMANCE_THRESHOLDS["generation_time_warning"]:
                logger.warning(
                    "Slow document generation",
                    template_name=template_name,
                    output_format=output_format,
                    generation_time=generation_time
                )
            
            logger.info(
                "Document generation completed",
                template_name=template_name,
                output_format=output_format,
                generation_time=generation_time,
                file_size=len(content),
                memory_delta=memory_delta
            )
            
            return {
                "content": content,
                "mime_type": mime_type,
                "metadata": {
                    "template_name": template_name,
                    "output_format": output_format,
                    "document_type": document_type,
                    "file_size": len(content),
                    "generation_time": generation_time,
                    "memory_usage_mb": memory_delta,
                    "timestamp": time.time()
                }
            }
            
        except Exception as e:
            generation_time = time.time() - start_time
            logger.error(
                "Document generation failed",
                template_name=template_name,
                output_format=output_format,
                error=str(e),
                generation_time=generation_time,
                exc_info=True
            )
            raise DocumentGenerationError(f"Document generation failed: {e}")
        
        finally:
            # Cleanup memory
            gc.collect()
    
    def _html_to_markdown(self, html_content: str) -> str:
        """Convert HTML to Markdown (simplified conversion)."""
        try:
            import re
            
            # Simple HTML to Markdown conversion
            # This is a basic implementation - could be enhanced with html2text library
            
            # Remove HTML tags we don't need in markdown
            markdown_content = re.sub(r'<!DOCTYPE[^>]*>', '', html_content)
            markdown_content = re.sub(r'<html[^>]*>|</html>', '', markdown_content)
            markdown_content = re.sub(r'<head>.*?</head>', '', markdown_content, flags=re.DOTALL)
            markdown_content = re.sub(r'<body[^>]*>|</body>', '', markdown_content)
            
            # Convert headings
            markdown_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', markdown_content)
            markdown_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', markdown_content)
            markdown_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', markdown_content)
            
            # Convert paragraphs
            markdown_content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', markdown_content)
            
            # Convert lists
            markdown_content = re.sub(r'<ul[^>]*>', '', markdown_content)
            markdown_content = re.sub(r'</ul>', '\n', markdown_content)
            markdown_content = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1', markdown_content)
            
            # Convert emphasis
            markdown_content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', markdown_content)
            markdown_content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', markdown_content)
            
            # Remove remaining HTML tags
            markdown_content = re.sub(r'<[^>]+>', '', markdown_content)
            
            # Clean up whitespace
            markdown_content = re.sub(r'\n\s*\n\s*\n', '\n\n', markdown_content)
            markdown_content = markdown_content.strip()
            
            return markdown_content
            
        except Exception as e:
            logger.error("HTML to Markdown conversion failed", error=str(e))
            # Fallback: return HTML with tags stripped
            import re
            return re.sub(r'<[^>]+>', '', html_content)
    
    async def validate_output_format(self, output_format: str) -> bool:
        """Validate if output format is supported."""
        supported_formats = self.settings.supported_formats
        return output_format.lower() in [fmt.lower() for fmt in supported_formats]
    
    async def get_template_info(self, template_name: str) -> Optional[Dict[str, Any]]:
        """Get information about a specific template."""
        try:
            # Load template configuration
            import yaml
            from ..core.config import DATA_DIR
            
            config_path = os.path.join(DATA_DIR, "template_configs.yaml")
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                
                templates = config.get('templates', {})
                return templates.get(template_name)
            
        except Exception as e:
            logger.error("Failed to load template info", template_name=template_name, error=str(e))
        
        return None
    
    async def list_available_templates(self, document_type: str = None) -> List[Dict[str, Any]]:
        """List all available templates, optionally filtered by document type."""
        try:
            import yaml
            from ..core.config import DATA_DIR
            
            config_path = os.path.join(DATA_DIR, "template_configs.yaml")
            if not os.path.exists(config_path):
                return []
            
            with open(config_path, 'r', encoding='utf-8') as f:
                config = yaml.safe_load(f)
            
            templates = config.get('templates', {})
            template_list = []
            
            for template_id, template_info in templates.items():
                if document_type is None or template_info.get('category') == document_type:
                    template_list.append({
                        'id': template_id,
                        'name': template_info.get('name', template_id),
                        'description': template_info.get('description', ''),
                        'category': template_info.get('category', 'unknown'),
                        'target_industries': template_info.get('target_industries', []),
                        'ats_optimization': template_info.get('ats_optimization', 'unknown')
                    })
            
            return template_list
            
        except Exception as e:
            logger.error("Failed to list templates", error=str(e))
            return []


# Memory monitoring and alerting
async def monitor_memory_usage() -> bool:
    """Monitor memory usage and trigger cleanup if needed."""
    process = psutil.Process()
    memory_percent = process.memory_percent()
    
    if memory_percent > 80:  # 80% threshold
        logger.warning(f"High memory usage: {memory_percent:.1f}%")
        
        # Trigger proactive cleanup
        gc.collect()
        
        if memory_percent > 90:  # Critical threshold
            logger.critical(f"Critical memory usage: {memory_percent:.1f}%")
            return False  # Signal to reject new requests
    
    return True  # Memory usage acceptable