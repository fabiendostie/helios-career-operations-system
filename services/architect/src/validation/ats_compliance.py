"""
ATS Compliance Validation System for ARCHITECT Service

This module provides comprehensive validation of generated documents against
2025 ATS (Applicant Tracking System) compliance standards including:

- **Format Compatibility**: PDF, DOCX, TXT parsing requirements
- **Content Structure**: Section headers, keyword density, formatting rules
- **Technical Standards**: Character encoding, font compatibility, file size limits
- **Parsing Optimization**: Machine readability, OCR compatibility, metadata standards

The system uses real-time intelligence from the research engine to stay current
with evolving ATS vendor requirements (Workday, Greenhouse, Lever, BambooHR, etc.).
"""

import re
import json
import asyncio
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass
from enum import Enum
import structlog
import time

# Document processing imports
from io import BytesIO
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup

logger = structlog.get_logger()

# Import research engine for real-time ATS requirements
try:
    from ..core.research_engine import get_research_engine
except ImportError:
    logger.warning("Research engine not available - using static fallbacks")
    get_research_engine = None

class ATSVendor(Enum):
    """Major ATS vendors with specific requirements."""
    WORKDAY = "workday"
    GREENHOUSE = "greenhouse"
    LEVER = "lever"
    BAMBOO_HR = "bamboo_hr"
    TALEO = "taleo"
    SUCCESSFACTORS = "successfactors"
    ICIMS = "icims"
    JOBVITE = "jobvite"
    GENERIC = "generic"

class ComplianceLevel(Enum):
    """ATS compliance validation levels."""
    STRICT = "strict"          # Maximum compatibility
    STANDARD = "standard"      # Industry standard
    BASIC = "basic"           # Minimal requirements

@dataclass
class ValidationResult:
    """Result of ATS compliance validation."""
    is_compliant: bool
    compliance_score: float  # 0-100
    violations: List[Dict[str, Any]]
    recommendations: List[str]
    ats_vendor_scores: Dict[ATSVendor, float]
    parsing_confidence: float
    # Enhanced fields for intelligent validation
    is_valid: Optional[bool] = None
    confidence_score: Optional[float] = None
    validation_issues: Optional[List['ValidationIssue']] = None
    vendor: Optional[ATSVendor] = None
    compliance_level: Optional[ComplianceLevel] = None

@dataclass
class ValidationIssue:
    """Individual validation issue found during ATS compliance check."""

    rule_id: str
    category: str
    severity: str
    is_valid: bool
    message: str
    recommendation: str

@dataclass
class ATSRule:
    """Individual ATS compliance rule."""
    rule_id: str
    description: str
    category: str
    severity: str  # critical, high, medium, low
    validator_func: callable
    applicable_vendors: List[ATSVendor]

class ATSComplianceValidator:
    """
    Comprehensive ATS compliance validation system.

    Validates documents against current ATS parsing requirements using
    real-time intelligence from the research engine.
    """

    def __init__(self):
        self.rules = []
        self.vendor_requirements = {}
        self.compliance_cache = {}
        self._initialize_validation_rules()

    def _initialize_validation_rules(self):
        """Initialize comprehensive ATS validation rules."""

        # Format and Structure Rules
        self.rules.extend([
            ATSRule(
                rule_id="format_pdf_text_extractable",
                description="PDF must have extractable text (not image-based)",
                category="format",
                severity="critical",
                validator_func=self._validate_pdf_text_extraction,
                applicable_vendors=[ATSVendor.WORKDAY, ATSVendor.GREENHOUSE, ATSVendor.LEVER]
            ),
            ATSRule(
                rule_id="format_file_size_limit",
                description="Document size must be under 2MB for optimal parsing",
                category="format",
                severity="high",
                validator_func=self._validate_file_size,
                applicable_vendors=list(ATSVendor)
            ),
            ATSRule(
                rule_id="format_font_compatibility",
                description="Use ATS-friendly fonts (Arial, Calibri, Times New Roman)",
                category="format",
                severity="medium",
                validator_func=self._validate_font_compatibility,
                applicable_vendors=[ATSVendor.TALEO, ATSVendor.SUCCESSFACTORS]
            ),
            ATSRule(
                rule_id="structure_section_headers",
                description="Clear section headers for Experience, Education, Skills",
                category="structure",
                severity="high",
                validator_func=self._validate_section_headers,
                applicable_vendors=list(ATSVendor)
            ),
            ATSRule(
                rule_id="structure_contact_info",
                description="Contact information at top with standard formatting",
                category="structure",
                severity="critical",
                validator_func=self._validate_contact_info,
                applicable_vendors=list(ATSVendor)
            ),
            ATSRule(
                rule_id="content_keyword_density",
                description="Appropriate keyword density for job relevance",
                category="content",
                severity="high",
                validator_func=self._validate_keyword_density,
                applicable_vendors=[ATSVendor.WORKDAY, ATSVendor.ICIMS]
            ),
            ATSRule(
                rule_id="content_date_formatting",
                description="Consistent date formatting (MM/YYYY or Month YYYY)",
                category="content",
                severity="medium",
                validator_func=self._validate_date_formatting,
                applicable_vendors=list(ATSVendor)
            ),
            ATSRule(
                rule_id="technical_encoding",
                description="UTF-8 encoding for special characters",
                category="technical",
                severity="medium",
                validator_func=self._validate_character_encoding,
                applicable_vendors=list(ATSVendor)
            ),
            ATSRule(
                rule_id="technical_no_headers_footers",
                description="Avoid complex headers/footers that confuse parsers",
                category="technical",
                severity="high",
                validator_func=self._validate_headers_footers,
                applicable_vendors=[ATSVendor.BAMBOO_HR, ATSVendor.JOBVITE]
            ),
            ATSRule(
                rule_id="technical_no_graphics_over_text",
                description="No graphics or images overlaying text content",
                category="technical",
                severity="critical",
                validator_func=self._validate_graphics_placement,
                applicable_vendors=list(ATSVendor)
            )
        ])

    async def validate_document(
        self,
        document_path: Path,
        target_vendors: List[ATSVendor] = None,
        compliance_level: ComplianceLevel = ComplianceLevel.STANDARD
    ) -> ValidationResult:
        """
        Perform comprehensive ATS compliance validation.

        Args:
            document_path: Path to document to validate
            target_vendors: Specific ATS vendors to optimize for
            compliance_level: Validation strictness level

        Returns:
            Detailed validation results with scores and recommendations
        """

        logger.info("Starting ATS compliance validation",
                   document=str(document_path),
                   target_vendors=[v.value if hasattr(v, 'value') else str(v) for v in (target_vendors or [])])

        if target_vendors is None:
            target_vendors = [ATSVendor.GENERIC]

        violations = []
        vendor_scores = {}

        # Extract document content for analysis
        try:
            content = await self._extract_document_content(document_path)
        except Exception as e:
            logger.error("Document content extraction failed", error=str(e))
            return ValidationResult(
                is_compliant=False,
                compliance_score=0.0,
                violations=[{
                    "rule_id": "extraction_failed",
                    "severity": "critical",
                    "description": f"Could not extract document content: {str(e)}"
                }],
                recommendations=["Ensure document is valid and accessible"],
                ats_vendor_scores={},
                parsing_confidence=0.0
            )

        # Run validation rules
        for rule in self.rules:
            # Skip rules not applicable to target vendors
            if not any(vendor in rule.applicable_vendors for vendor in target_vendors):
                continue

            try:
                is_valid, details = await rule.validator_func(content, document_path)

                if not is_valid:
                    violations.append({
                        "rule_id": rule.rule_id,
                        "description": rule.description,
                        "category": rule.category,
                        "severity": rule.severity,
                        "details": details,
                        "applicable_vendors": [v.value for v in rule.applicable_vendors]
                    })

            except Exception as e:
                logger.warning("Validation rule failed",
                             rule_id=rule.rule_id, error=str(e))
                violations.append({
                    "rule_id": rule.rule_id,
                    "description": f"Validation error: {str(e)}",
                    "category": "system",
                    "severity": "medium",
                    "details": {},
                    "applicable_vendors": []
                })

        # Calculate vendor-specific scores
        for vendor in target_vendors:
            vendor_violations = [v for v in violations
                               if vendor.value in v.get("applicable_vendors", [])]
            vendor_scores[vendor] = self._calculate_vendor_score(vendor_violations)

        # Calculate overall compliance score
        compliance_score = self._calculate_overall_score(violations)
        parsing_confidence = self._estimate_parsing_confidence(content, violations)

        # Generate recommendations
        recommendations = self._generate_recommendations(violations, target_vendors)

        is_compliant = compliance_score >= self._get_compliance_threshold(compliance_level)

        logger.info("ATS compliance validation completed",
                   compliance_score=compliance_score,
                   violations_count=len(violations),
                   is_compliant=is_compliant)

        return ValidationResult(
            is_compliant=is_compliant,
            compliance_score=compliance_score,
            violations=violations,
            recommendations=recommendations,
            ats_vendor_scores=vendor_scores,
            parsing_confidence=parsing_confidence
        )

    async def _extract_document_content(self, document_path: Path) -> Dict[str, Any]:
        """Extract content and metadata from document for analysis."""

        content = {
            "file_path": str(document_path),
            "file_size": document_path.stat().st_size,
            "file_extension": document_path.suffix.lower(),
            "raw_text": "",
            "structured_content": {},
            "metadata": {}
        }

        if document_path.suffix.lower() == '.pdf':
            content.update(await self._extract_pdf_content(document_path))
        elif document_path.suffix.lower() in ['.docx', '.doc']:
            content.update(await self._extract_docx_content(document_path))
        elif document_path.suffix.lower() in ['.html', '.htm']:
            content.update(await self._extract_html_content(document_path))
        elif document_path.suffix.lower() == '.txt':
            content.update(await self._extract_txt_content(document_path))
        else:
            raise ValueError(f"Unsupported document format: {document_path.suffix}")

        return content

    async def _extract_pdf_content(self, pdf_path: Path) -> Dict[str, Any]:
        """Extract content from PDF document."""

        content = {"raw_text": "", "metadata": {}, "extraction_method": "pdf"}

        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata
                if pdf_reader.metadata:
                    content["metadata"] = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "creator": pdf_reader.metadata.get('/Creator', '')
                    }

                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning("PDF page extraction failed",
                                     page=page_num, error=str(e))

                content["raw_text"] = "\n".join(text_content)
                content["page_count"] = len(pdf_reader.pages)

        except Exception as e:
            logger.error("PDF content extraction failed", error=str(e))
            content["extraction_error"] = str(e)

        return content

    async def _extract_docx_content(self, docx_path: Path) -> Dict[str, Any]:
        """Extract content from DOCX document."""

        content = {"raw_text": "", "metadata": {}, "extraction_method": "docx"}

        try:
            doc = Document(docx_path)

            # Extract core properties
            content["metadata"] = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else ""
            }

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)

            content["raw_text"] = "\n".join(paragraphs)

            # Extract font information for validation
            fonts_used = set()
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.name:
                        fonts_used.add(run.font.name)

            content["fonts_used"] = list(fonts_used)

        except Exception as e:
            logger.error("DOCX content extraction failed", error=str(e))
            content["extraction_error"] = str(e)

        return content

    async def _extract_html_content(self, html_path: Path) -> Dict[str, Any]:
        """Extract content from HTML document."""

        content = {"raw_text": "", "metadata": {}, "extraction_method": "html"}

        try:
            with open(html_path, 'r', encoding='utf-8') as file:
                html_content = file.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Extract metadata
            title_tag = soup.find('title')
            content["metadata"]["title"] = title_tag.text if title_tag else ""

            # Extract text content
            content["raw_text"] = soup.get_text(separator='\n', strip=True)

            # Check for CSS styling that might affect ATS parsing
            style_tags = soup.find_all('style')
            content["has_complex_styling"] = len(style_tags) > 0

        except Exception as e:
            logger.error("HTML content extraction failed", error=str(e))
            content["extraction_error"] = str(e)

        return content

    async def _extract_txt_content(self, txt_path: Path) -> Dict[str, Any]:
        """Extract content from plain text document."""

        content = {"raw_text": "", "metadata": {}, "extraction_method": "txt"}

        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                content["raw_text"] = file.read()

        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(txt_path, 'r', encoding=encoding) as file:
                        content["raw_text"] = file.read()
                        content["encoding_used"] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                content["extraction_error"] = "Could not decode text file"

        except Exception as e:
            logger.error("TXT content extraction failed", error=str(e))
            content["extraction_error"] = str(e)

        return content

    # Validation Rule Implementations

    async def _validate_pdf_text_extraction(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate PDF has extractable text content."""

        if content["file_extension"] != ".pdf":
            return True, {}  # Not applicable

        raw_text = content.get("raw_text", "")

        # Check if text was successfully extracted
        if not raw_text or len(raw_text.strip()) < 50:
            return False, {
                "extracted_chars": len(raw_text),
                "issue": "Insufficient extractable text - may be image-based PDF"
            }

        # Check for signs of OCR text (lots of spacing/formatting issues)
        ocr_indicators = [
            len(re.findall(r'\s{3,}', raw_text)) > 10,  # Excessive spacing
            len(re.findall(r'\n\s*\n', raw_text)) / len(raw_text) > 0.05,  # Too many line breaks
        ]

        if any(ocr_indicators):
            return False, {
                "issue": "Text appears to be from OCR - may have parsing issues"
            }

        return True, {"extracted_chars": len(raw_text)}

    async def _validate_file_size(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate document file size is within ATS limits."""

        file_size = content.get("file_size", 0)
        max_size = 2 * 1024 * 1024  # 2MB

        if file_size > max_size:
            return False, {
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "max_size_mb": 2,
                "issue": "File size exceeds ATS processing limits"
            }

        return True, {"file_size_mb": round(file_size / (1024 * 1024), 2)}

    async def _validate_font_compatibility(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate use of ATS-friendly fonts."""

        # Only applicable to DOCX files where font info is available
        fonts_used = content.get("fonts_used", [])
        if not fonts_used:
            return True, {}  # Cannot validate, assume OK

        ats_friendly_fonts = {
            "Arial", "Calibri", "Times New Roman", "Helvetica",
            "Georgia", "Verdana", "Tahoma", "Trebuchet MS"
        }

        problematic_fonts = [font for font in fonts_used if font not in ats_friendly_fonts]

        if problematic_fonts:
            return False, {
                "problematic_fonts": problematic_fonts,
                "recommended_fonts": list(ats_friendly_fonts)
            }

        return True, {"fonts_used": fonts_used}

    async def _validate_section_headers(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate presence of clear section headers."""

        raw_text = content.get("raw_text", "").lower()

        required_sections = {
            "contact": ["contact", "personal information"],
            "experience": ["experience", "work experience", "employment", "professional experience"],
            "education": ["education", "academic background", "qualifications"],
            "skills": ["skills", "technical skills", "competencies"]
        }

        missing_sections = []
        found_sections = {}

        for section_type, keywords in required_sections.items():
            found = any(keyword in raw_text for keyword in keywords)
            found_sections[section_type] = found
            if not found:
                missing_sections.append(section_type)

        if missing_sections:
            return False, {
                "missing_sections": missing_sections,
                "found_sections": found_sections,
                "issue": "Missing standard resume sections"
            }

        return True, {"found_sections": found_sections}

    async def _validate_contact_info(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate contact information is present and well-formatted."""

        raw_text = content.get("raw_text", "")

        # Check for email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        has_email = bool(re.search(email_pattern, raw_text))

        # Check for phone number
        phone_pattern = r'(\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        has_phone = bool(re.search(phone_pattern, raw_text))

        issues = []
        if not has_email:
            issues.append("No email address found")
        if not has_phone:
            issues.append("No phone number found")

        if issues:
            return False, {
                "has_email": has_email,
                "has_phone": has_phone,
                "issues": issues
            }

        return True, {"has_email": has_email, "has_phone": has_phone}

    async def _validate_keyword_density(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate appropriate keyword density for ATS optimization."""

        raw_text = content.get("raw_text", "")

        if not raw_text:
            return False, {"issue": "No text content to analyze"}

        # Simple keyword density analysis
        words = raw_text.lower().split()
        word_count = len(words)

        # Check for keyword stuffing (repeated words)
        word_frequency = {}
        for word in words:
            if len(word) > 3:  # Only count meaningful words
                word_frequency[word] = word_frequency.get(word, 0) + 1

        # Find words that appear too frequently
        stuffed_words = []
        for word, count in word_frequency.items():
            frequency = count / word_count
            if frequency > 0.02 and count > 5:  # More than 2% frequency
                stuffed_words.append((word, count, f"{frequency:.1%}"))

        if stuffed_words:
            return False, {
                "stuffed_words": stuffed_words[:5],  # Top 5 offenders
                "issue": "Potential keyword stuffing detected"
            }

        return True, {"word_count": word_count}

    async def _validate_date_formatting(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate consistent date formatting."""

        raw_text = content.get("raw_text", "")

        # Find various date patterns
        date_patterns = [
            (r'\b\d{1,2}/\d{4}\b', 'MM/YYYY'),
            (r'\b\d{4}\b', 'YYYY'),
            (r'\b[A-Za-z]{3,9}\s+\d{4}\b', 'Month YYYY'),
            (r'\b\d{1,2}/\d{1,2}/\d{4}\b', 'MM/DD/YYYY'),
        ]

        found_formats = []
        for pattern, format_name in date_patterns:
            matches = re.findall(pattern, raw_text)
            if matches:
                found_formats.append((format_name, len(matches)))

        # Check for consistency - should primarily use one format
        if len(found_formats) > 2:
            return False, {
                "found_formats": found_formats,
                "issue": "Inconsistent date formatting may confuse ATS parsers"
            }

        return True, {"date_formats": found_formats}

    async def _validate_character_encoding(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate proper character encoding."""

        raw_text = content.get("raw_text", "")

        # Check for problematic characters that might not parse well
        problematic_chars = []
        for char in raw_text:
            if ord(char) > 127 and char not in '""''—–':  # Allow common smart quotes/dashes
                problematic_chars.append(char)

        # Count unique problematic characters
        unique_problematic = list(set(problematic_chars))

        if len(unique_problematic) > 5:
            return False, {
                "problematic_chars": unique_problematic[:10],
                "count": len(problematic_chars),
                "issue": "Many non-ASCII characters may cause parsing issues"
            }

        return True, {"non_ascii_chars": len(unique_problematic)}

    async def _validate_headers_footers(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate minimal use of headers/footers."""

        # This is a simplified check - in practice would need format-specific analysis
        raw_text = content.get("raw_text", "")

        # Look for repeated content that might indicate headers/footers
        lines = [line.strip() for line in raw_text.split('\n') if line.strip()]

        if len(lines) < 5:
            return True, {}  # Too short to analyze

        # Check if first/last lines repeat frequently (header/footer pattern)
        first_line = lines[0] if lines else ""
        last_line = lines[-1] if lines else ""

        # Simple heuristic: if first/last line contains only page numbers or dates
        header_footer_patterns = [
            r'^\s*\d+\s*$',  # Just page numbers
            r'^\s*page\s+\d+\s*$',  # "Page X"
            r'^\s*\d{1,2}/\d{1,2}/\d{4}\s*$',  # Just dates
        ]

        has_header_footer = any(
            re.match(pattern, first_line, re.IGNORECASE) or
            re.match(pattern, last_line, re.IGNORECASE)
            for pattern in header_footer_patterns
        )

        if has_header_footer:
            return False, {
                "first_line": first_line,
                "last_line": last_line,
                "issue": "Detected headers/footers that may interfere with parsing"
            }

        return True, {}

    async def _validate_graphics_placement(self, content: Dict[str, Any], document_path: Path) -> Tuple[bool, Dict]:
        """Validate no graphics overlay text content."""

        # This is a simplified check - full implementation would require image analysis
        raw_text = content.get("raw_text", "")

        # Heuristic: check for very sparse text that might indicate text-over-images
        if not raw_text:
            return False, {"issue": "No extractable text found"}

        lines = [line.strip() for line in raw_text.split('\n') if line.strip()]
        total_chars = sum(len(line) for line in lines)
        avg_line_length = total_chars / len(lines) if lines else 0

        # If average line length is very short, might indicate parsing issues from graphics
        if avg_line_length < 20 and len(lines) > 10:
            return False, {
                "avg_line_length": avg_line_length,
                "line_count": len(lines),
                "issue": "Very short text lines may indicate graphics interfering with text"
            }

        return True, {"avg_line_length": avg_line_length}

    # Scoring and Analysis Methods

    def _calculate_vendor_score(self, violations: List[Dict]) -> float:
        """Calculate compliance score for specific ATS vendor."""

        if not violations:
            return 100.0

        severity_weights = {
            "critical": 25,
            "high": 15,
            "medium": 10,
            "low": 5
        }

        total_penalty = sum(severity_weights.get(v["severity"], 5) for v in violations)
        return max(0.0, 100.0 - total_penalty)

    def _calculate_overall_score(self, violations: List[Dict]) -> float:
        """Calculate overall compliance score."""

        if not violations:
            return 100.0

        severity_weights = {
            "critical": 20,
            "high": 12,
            "medium": 8,
            "low": 4
        }

        total_penalty = sum(severity_weights.get(v["severity"], 4) for v in violations)
        return max(0.0, 100.0 - total_penalty)

    def _estimate_parsing_confidence(self, content: Dict[str, Any], violations: List[Dict]) -> float:
        """Estimate ATS parsing confidence based on content and violations."""

        base_confidence = 90.0

        # Reduce confidence based on critical issues
        critical_violations = [v for v in violations if v.get("severity") == "critical"]
        base_confidence -= len(critical_violations) * 20

        # Reduce based on extraction quality
        raw_text = content.get("raw_text", "") or ""  # Handle None values
        if len(raw_text) < 100:
            base_confidence -= 30
        elif "extraction_error" in content:
            base_confidence -= 40

        return max(0.0, min(100.0, base_confidence))

    def _generate_recommendations(self, violations: List[Dict], target_vendors: List[ATSVendor]) -> List[str]:
        """Generate actionable recommendations based on violations."""

        recommendations = []

        # Group violations by category
        by_category = {}
        for violation in violations:
            category = violation.get("category", "general")
            if category not in by_category:
                by_category[category] = []
            by_category[category].append(violation)

        # Generate category-specific recommendations
        if "format" in by_category:
            recommendations.append("Consider converting to simpler PDF format or DOCX for better ATS compatibility")

        if "structure" in by_category:
            recommendations.append("Use clear section headers: Contact Information, Professional Experience, Education, Skills")

        if "content" in by_category:
            recommendations.append("Review keyword usage - avoid repetition while maintaining job relevance")

        if "technical" in by_category:
            recommendations.append("Use standard fonts (Arial, Calibri) and avoid complex formatting")

        # Add vendor-specific recommendations
        if ATSVendor.WORKDAY in target_vendors:
            recommendations.append("Workday systems prefer chronological format with clear date ranges")

        if ATSVendor.GREENHOUSE in target_vendors:
            recommendations.append("Greenhouse parses skills sections best when formatted as bulleted lists")

        return recommendations

    async def _get_current_ats_requirements(self) -> Dict[str, Any]:
        """Get current ATS requirements from internet research."""
        logger.info("Fetching current ATS requirements from internet research")

        # Try to get real-time intelligence from research engine
        if get_research_engine:
            try:
                research_engine = await get_research_engine()
                ats_intelligence = await research_engine.get_ats_compliance_intelligence()

                logger.info("Successfully retrieved current ATS requirements",
                           vendor_count=len(ats_intelligence.get('vendor_requirements', {})),
                           last_updated=ats_intelligence.get('last_updated'))

                return ats_intelligence

            except Exception as e:
                logger.warning("Failed to retrieve real-time ATS requirements", error=str(e))

        # Fallback to static requirements with current date
        logger.info("Using fallback static ATS requirements")
        return {
            'vendor_requirements': {
                'workday': {
                    'supported_formats': ['.pdf', '.docx'],
                    'max_file_size_mb': 5,
                    'required_sections': ['contact', 'experience', 'education'],
                    'font_compatibility': ['Arial', 'Times New Roman', 'Calibri'],
                    'parsing_preferences': {'keyword_density_max': 0.15, 'section_order_strict': False}
                },
                'greenhouse': {
                    'supported_formats': ['.pdf', '.docx', '.txt'],
                    'max_file_size_mb': 10,
                    'required_sections': ['contact', 'experience'],
                    'font_compatibility': ['Arial', 'Helvetica', 'Times New Roman'],
                    'parsing_preferences': {'keyword_density_max': 0.20, 'section_order_strict': True}
                },
                'lever': {
                    'supported_formats': ['.pdf', '.docx'],
                    'max_file_size_mb': 8,
                    'required_sections': ['contact', 'experience', 'skills'],
                    'font_compatibility': ['Arial', 'Calibri', 'Helvetica'],
                    'parsing_preferences': {'keyword_density_max': 0.18, 'skills_section_required': True}
                }
            },
            'updated_standards': {
                'pdf_version_min': '1.4',
                'encoding_required': 'UTF-8',
                'accessibility_compliance': 'WCAG_2.1',
                'metadata_requirements': ['title', 'author', 'creation_date']
            },
            'current_trends': {
                'ai_parsing_compatibility': True,
                'semantic_analysis_support': True,
                'multilingual_support': ['en', 'es', 'fr']
            },
            'last_updated': time.strftime('%Y-%m-%dT%H:%M:%SZ'),
            'source': 'static_fallback'
        }

    async def validate_for_vendor(self, content: Dict[str, Any], file_path: Path, vendor: ATSVendor) -> 'ValidationResult':
        """Validate document against specific ATS vendor requirements."""
        logger.info("Validating for specific ATS vendor", vendor=vendor.value)

        # Get current requirements
        current_requirements = await self._get_current_ats_requirements()
        vendor_reqs = current_requirements.get('vendor_requirements', {}).get(vendor.value, {})

        validation_issues = []

        # Validate format compatibility
        file_ext = file_path.suffix.lower()
        supported_formats = vendor_reqs.get('supported_formats', ['.pdf', '.docx'])
        if file_ext not in supported_formats:
            validation_issues.append(ValidationIssue(
                rule_id='vendor_format_compatibility',
                category='format',
                severity='high',
                is_valid=False,
                message=f"{vendor.value} does not support {file_ext} format. Supported: {', '.join(supported_formats)}",
                recommendation=f"Convert document to {supported_formats[0]} format for {vendor.value} compatibility"
            ))

        # Validate file size
        file_size_mb = content.get('file_size', 0) / (1024 * 1024)
        max_size = vendor_reqs.get('max_file_size_mb', 10)
        if file_size_mb > max_size:
            validation_issues.append(ValidationIssue(
                rule_id='vendor_file_size',
                category='technical',
                severity='critical',
                is_valid=False,
                message=f"File size {file_size_mb:.1f}MB exceeds {vendor.value} limit of {max_size}MB",
                recommendation=f"Reduce file size to under {max_size}MB for {vendor.value}"
            ))

        # Validate required sections
        required_sections = vendor_reqs.get('required_sections', ['contact', 'experience'])
        content_sections = content.get('sections', {})
        for section in required_sections:
            if section not in content_sections or not content_sections[section]:
                validation_issues.append(ValidationIssue(
                    rule_id=f'vendor_section_{section}',
                    category='structure',
                    severity='medium',
                    is_valid=False,
                    message=f"{vendor.value} requires '{section}' section",
                    recommendation=f"Add clear '{section}' section header and content"
                ))

        is_valid = all(issue.is_valid for issue in validation_issues)
        violations_dict = [v.__dict__ for v in validation_issues]
        confidence_score = self._calculate_confidence(content, violations_dict)

        return ValidationResult(
            is_compliant=is_valid,
            compliance_score=confidence_score,
            violations=violations_dict,
            recommendations=self._generate_recommendations(violations_dict, [vendor]),
            ats_vendor_scores={vendor: confidence_score},
            parsing_confidence=confidence_score,
            is_valid=is_valid,
            confidence_score=confidence_score,
            validation_issues=validation_issues,
            vendor=vendor,
            compliance_level=ComplianceLevel.STANDARD
        )

    async def _validate_current_standards(self, content: Dict[str, Any], file_path: Path) -> Tuple[bool, Dict[str, Any]]:
        """Validate against current parsing standards from internet research."""
        current_requirements = await self._get_current_ats_requirements()
        standards = current_requirements.get('updated_standards', {})

        validation_details = {'validation_details': {}}
        issues = []

        # Validate encoding
        metadata = content.get('metadata', {})
        required_encoding = standards.get('encoding_required', 'UTF-8')
        actual_encoding = metadata.get('encoding', 'unknown')

        validation_details['validation_details']['encoding'] = {
            'required': required_encoding,
            'actual': actual_encoding,
            'valid': actual_encoding == required_encoding
        }

        if actual_encoding != required_encoding:
            issues.append(f"Encoding should be {required_encoding}, found {actual_encoding}")

        # Check AI compatibility
        ai_compatible = current_requirements.get('current_trends', {}).get('ai_parsing_compatibility', False)
        validation_details['validation_details']['ai_compatibility'] = {
            'supported': ai_compatible,
            'text_quality_score': len(content.get('raw_text', '')) / max(1, content.get('word_count', 1))
        }

        # Validate metadata completeness
        required_metadata = standards.get('metadata_requirements', [])
        metadata_completeness = {}
        for field in required_metadata:
            present = field in metadata and metadata[field]
            metadata_completeness[field] = present
            if not present:
                issues.append(f"Missing required metadata field: {field}")

        validation_details['validation_details']['metadata_completeness'] = metadata_completeness

        # Assess semantic structure
        text = content.get('raw_text', '')
        sections = content.get('sections', {})
        validation_details['validation_details']['semantic_structure'] = {
            'clear_sections': len(sections) >= 3,
            'adequate_content': len(text) >= 200,
            'structured_format': bool(sections)
        }

        is_valid = len(issues) == 0
        return is_valid, validation_details

    async def _analyze_content_structure(self, content: Dict[str, Any], file_path: Path) -> Dict[str, Any]:
        """Intelligent analysis of content structure for ATS optimization."""
        text = content.get('raw_text', '')
        sections = content.get('sections', {})

        analysis = {}

        # Section completeness analysis
        section_completeness = {
            'contact_info_complete': bool(sections.get('contact') and '@' in sections.get('contact', '')),
            'experience_detailed': bool(sections.get('experience') and len(sections.get('experience', '')) > 100),
            'skills_categorized': bool(sections.get('skills') and len(sections.get('skills', '').split(',')) > 3),
            'education_present': bool(sections.get('education'))
        }
        analysis['section_completeness'] = section_completeness

        # Keyword analysis
        words = text.lower().split()
        technical_keywords = ['python', 'javascript', 'aws', 'docker', 'kubernetes', 'react', 'java', 'sql']
        action_verbs = ['led', 'managed', 'developed', 'implemented', 'designed', 'optimized', 'achieved']

        keyword_analysis = {
            'technical_keywords_count': sum(1 for word in words if word in technical_keywords),
            'action_verbs_count': sum(1 for word in words if word in action_verbs),
            'quantified_achievements_count': len(re.findall(r'\d+%|\$\d+|x\d+|\d+,\d+', text))
        }
        analysis['keyword_analysis'] = keyword_analysis

        # Parser readability assessment
        parser_readability = {
            'text_extraction_quality': min(1.0, len(text) / max(1, content.get('character_count', 1))),
            'formatting_consistency': len(sections) >= 3,
            'section_hierarchy_clear': bool(re.search(r'(EXPERIENCE|EDUCATION|SKILLS)', text.upper()))
        }
        analysis['parser_readability'] = parser_readability

        # Optimization suggestions
        suggestions = []
        if keyword_analysis['quantified_achievements_count'] < 3:
            suggestions.append("Add more quantified achievements (percentages, dollar amounts, metrics)")
        if not section_completeness['skills_categorized']:
            suggestions.append("Organize skills into clear categories (Technical Skills, Soft Skills)")
        if keyword_analysis['action_verbs_count'] < 5:
            suggestions.append("Use more strong action verbs to describe accomplishments")

        analysis['optimization_suggestions'] = suggestions

        return analysis

    async def validate_with_adaptive_compliance(self, content: Dict[str, Any], file_path: Path,
                                              industry: str, target_vendors: List[ATSVendor]) -> 'ValidationResult':
        """Validate with adaptive compliance levels based on current market standards."""
        current_requirements = await self._get_current_ats_requirements()
        market_intel = current_requirements.get('market_intelligence', {})

        # Determine compliance level based on industry
        strict_industries = market_intel.get('strict_compliance_industries', ['finance', 'healthcare', 'government'])
        compliance_level = ComplianceLevel.STRICT if industry.lower() in strict_industries else ComplianceLevel.STANDARD

        logger.info("Adaptive compliance validation", industry=industry,
                   compliance_level=compliance_level.value, vendor_count=len(target_vendors))

        # Run validation with determined compliance level
        validation_issues = []

        # Apply stricter rules for strict compliance
        if compliance_level == ComplianceLevel.STRICT:
            # Additional strict validation rules
            metadata = content.get('metadata', {})
            if not metadata.get('title'):
                validation_issues.append(ValidationIssue(
                    rule_id='strict_metadata_title',
                    category='technical',
                    severity='high',
                    is_valid=False,
                    message="Document title metadata required for strict compliance",
                    recommendation="Add document title in metadata"
                ))

        # Multi-vendor validation
        for vendor in target_vendors:
            vendor_result = await self.validate_for_vendor(content, file_path, vendor)
            validation_issues.extend(vendor_result.validation_issues)

        is_valid = len([i for i in validation_issues if not i.is_valid]) == 0
        violations_dict = [v.__dict__ for v in validation_issues]
        confidence_score = self._calculate_confidence(content, violations_dict)

        return ValidationResult(
            is_compliant=is_valid,
            compliance_score=confidence_score,
            violations=violations_dict,
            recommendations=self._generate_recommendations(violations_dict, target_vendors),
            ats_vendor_scores={v: confidence_score for v in target_vendors},
            parsing_confidence=confidence_score,
            is_valid=is_valid,
            confidence_score=confidence_score,
            validation_issues=validation_issues,
            vendor=target_vendors[0] if target_vendors else ATSVendor.GENERIC,
            compliance_level=compliance_level
        )

    async def _analyze_document_parseability(self, content: Dict[str, Any], file_path: Path) -> Dict[str, Any]:
        """Analyze document parsing difficulty and ATS compatibility."""
        analysis = {}

        # Calculate parsing difficulty score
        complexity_factors = []

        # Check formatting complexity
        if content.get('formatting_complexity') == 'high':
            complexity_factors.append(0.3)

        # Check for embedded objects
        embedded_objects = content.get('structured_content', {}).get('embedded_objects', [])
        if embedded_objects:
            complexity_factors.append(0.2)

        # Check text extraction quality
        raw_text_length = len(content.get('raw_text', ''))
        expected_length = content.get('character_count', raw_text_length)
        if expected_length > 0:
            extraction_quality = raw_text_length / expected_length
            if extraction_quality < 0.8:
                complexity_factors.append(0.4)

        parsing_difficulty = min(1.0, sum(complexity_factors))
        analysis['parsing_difficulty'] = parsing_difficulty

        # Calculate ATS compatibility score
        ats_compatibility = 1.0 - parsing_difficulty
        analysis['ats_compatibility_score'] = ats_compatibility

        # Assess structure quality
        sections = content.get('sections', {})
        structure_quality = min(1.0, len(sections) / 5.0)  # Expect 5 main sections
        analysis['structure_quality'] = structure_quality

        # Generate optimization recommendations
        recommendations = []
        if parsing_difficulty > 0.5:
            recommendations.append("Simplify document formatting for better ATS parsing")
        if embedded_objects:
            recommendations.append("Remove embedded objects that may interfere with text extraction")
        if structure_quality < 0.6:
            recommendations.append("Add clear section headers for better content organization")

        analysis['optimization_recommendations'] = recommendations

        return analysis

    async def _assess_semantic_content_quality(self, content: Dict[str, Any], file_path: Path) -> Dict[str, Any]:
        """Assess semantic quality of resume content for ATS systems."""
        text = content.get('raw_text', '')
        sections = content.get('sections', {})

        analysis = {}

        # Content depth analysis
        word_count = len(text.split())
        unique_words = len(set(text.lower().split()))
        content_depth = min(1.0, unique_words / max(1, word_count * 0.7))  # Expect 70% unique words
        analysis['content_depth_score'] = content_depth

        # Skill context quality
        skills_text = sections.get('skills', '')
        experience_text = sections.get('experience', '')

        # Check for skills mentioned in context
        skills_list = [skill.strip() for skill in skills_text.replace(',', ' ').split() if len(skill.strip()) > 2]
        skills_in_context = sum(1 for skill in skills_list if skill.lower() in experience_text.lower())

        skill_context_quality = {
            'technical_skills_with_context': skills_in_context,
            'experience_alignment_score': skills_in_context / max(1, len(skills_list))
        }
        analysis['skill_context_quality'] = skill_context_quality

        # Achievement quality assessment
        quantified_pattern = r'\d+%|\$\d+|x\d+|\d+,\d+|\d+ [a-zA-Z]+'
        quantified_results = re.findall(quantified_pattern, text)

        action_verbs = ['achieved', 'improved', 'increased', 'reduced', 'delivered', 'led', 'managed', 'developed']
        action_statements = sum(1 for verb in action_verbs if verb in text.lower())

        achievement_quality = {
            'quantified_results_count': len(quantified_results),
            'impact_statements_quality': action_statements / max(1, word_count / 100),  # Per 100 words
            'action_verb_strength': min(1.0, action_statements / 10)  # Expect ~10 action verbs
        }
        analysis['achievement_quality'] = achievement_quality

        # Content enhancement suggestions
        suggestions = []
        if content_depth < 0.6:
            suggestions.append("Expand content with more detailed descriptions and specific examples")
        if skill_context_quality['experience_alignment_score'] < 0.5:
            suggestions.append("Provide context for technical skills by mentioning them in experience descriptions")
        if achievement_quality['quantified_results_count'] < 3:
            suggestions.append("Add quantified achievements with specific metrics and percentages")
        if achievement_quality['action_verb_strength'] < 0.5:
            suggestions.append("Use stronger action verbs to describe accomplishments and responsibilities")

        analysis['enhancement_suggestions'] = suggestions

        return analysis

    async def analyze_multi_vendor_compatibility(self, content: Dict[str, Any], file_path: Path,
                                               target_vendors: List[ATSVendor]) -> Dict[str, Any]:
        """Analyze compatibility across multiple ATS vendors simultaneously."""
        logger.info("Analyzing multi-vendor compatibility", vendor_count=len(target_vendors))

        results = {}
        vendor_scores = {}
        all_issues = []
        vendor_recommendations = {}

        # Validate against each vendor
        for vendor in target_vendors:
            vendor_result = await self.validate_for_vendor(content, file_path, vendor)
            vendor_scores[vendor.value] = vendor_result.confidence_score / 100.0
            all_issues.extend(vendor_result.validation_issues)
            vendor_recommendations[vendor.value] = vendor_result.recommendations

        results['vendor_compatibility_scores'] = vendor_scores

        # Identify common issues across vendors
        issue_counts = {}
        for issue in all_issues:
            rule_id = issue.rule_id
            issue_counts[rule_id] = issue_counts.get(rule_id, 0) + 1

        common_issues = [
            rule_id for rule_id, count in issue_counts.items()
            if count >= len(target_vendors) * 0.5  # Issues affecting 50%+ of vendors
        ]
        results['common_issues'] = common_issues

        results['vendor_specific_recommendations'] = vendor_recommendations

        # Overall compatibility score
        overall_score = sum(vendor_scores.values()) / len(vendor_scores) if vendor_scores else 0
        results['overall_compatibility_score'] = overall_score

        return results

    def _calculate_confidence(self, content: Dict[str, Any], violations: List[Dict]) -> float:
        """Calculate confidence score based on content quality and violations."""
        base_confidence = 95.0

        # Reduce confidence based on violations
        for violation in violations:
            severity = violation.get('severity', 'medium')
            if severity == 'critical':
                base_confidence -= 15
            elif severity == 'high':
                base_confidence -= 10
            elif severity == 'medium':
                base_confidence -= 5
            else:  # low
                base_confidence -= 2

        # Adjust based on content quality
        raw_text = content.get('raw_text', '')
        if len(raw_text) < 100:
            base_confidence -= 20
        elif len(raw_text) < 500:
            base_confidence -= 10

        # Check for extraction errors
        if 'extraction_error' in content:
            base_confidence -= 25

        return max(0.0, min(100.0, base_confidence))

    def _get_compliance_threshold(self, level: ComplianceLevel) -> float:
        """Get compliance score threshold for different levels."""

        thresholds = {
            ComplianceLevel.STRICT: 95.0,
            ComplianceLevel.STANDARD: 85.0,
            ComplianceLevel.BASIC: 70.0
        }

        return thresholds.get(level, 85.0)

# Factory function for easy instantiation
async def get_ats_validator() -> ATSComplianceValidator:
    """Get configured ATS compliance validator instance."""

    validator = ATSComplianceValidator()

    # Load any cached vendor requirements or real-time updates
    # (This would integrate with the research engine for live ATS updates)

    return validator
